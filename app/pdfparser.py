import os
from typing import List, Union
import re
import tqdm

from langchain.schema import Document
# import spacy
import PyPDF2


def extract_page_text(filepath, max_len=512, overlap_len=100):
    page_content  = []
    # spliter = spacy.load("zh_core_web_sm")
    # chunks = []
    with open(filepath, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        page_count = 0
        # pattern = r'^\d{1,3}'
        for page in tqdm.tqdm(pdf_reader.pages):
            page_text = page.extract_text().strip()
            raw_text = [text.strip() for text in page_text.split('\n')]
            new_text = '\n'.join(raw_text)
            new_text = re.sub(r'\n\d{2,3}\s?', '\n', new_text) # remove page number
            # new_text = re.sub(pattern, '', new_text).strip()
            if len(new_text)>10 and '..............' not in new_text:
                page_content.append(new_text)


    cleaned_chunks = []
    i = 0
    all_str = ''.join(page_content)
    all_str = all_str.replace('\n', '')
    while i<len(all_str):
        cur_s = all_str[i:i+max_len]
        if len(cur_s)>10:
            cleaned_chunks.append(Document(page_content=cur_s, metadata={'page':page_count+1}))
        i+=(max_len - overlap_len)

    return cleaned_chunks

class FileParser:
    def __init__(self, files: Union[str, List[str]] = None):
        """Initialize the FileParser class."""
        self.corpus = ''
        if files:
            self.corpus = self.add_corpus(files)

    def add_corpus(self, files: Union[str, List[str]]):
        """Load document files."""
        if isinstance(files, str):
            files = [files]
        for doc_file in files:
            if doc_file.endswith('.pdf'):
                corpus = self.extract_text_from_pdf(doc_file)
            elif doc_file.endswith('.docx'):
                corpus = self.extract_text_from_docx(doc_file)
            elif doc_file.endswith('.md'):
                corpus = self.extract_text_from_markdown(doc_file)
            else:
                corpus = self.extract_text_from_txt(doc_file)
            full_text = '\n'.join(corpus)
        return full_text
            
    @staticmethod
    def extract_text_from_pdf(file_path: str):
        """Extract text content from a PDF file."""
        import PyPDF2
        contents = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text().strip()
                raw_text = [text.strip() for text in page_text.splitlines() if text.strip()]
                new_text = ''
                for text in raw_text:
                    new_text += text
                    if text[-1] in ['.', '!', '?', '。', '！', '？', '…', ';', '；', ':', '：', '”', '’', '）', '】', '》', '」',
                                    '』', '〕', '〉', '》', '〗', '〞', '〟', '»', '"', "'", ')', ']', '}']:
                        contents.append(new_text)
                        new_text = ''
                if new_text:
                    contents.append(new_text)
        return contents

    @staticmethod
    def extract_text_from_txt(file_path: str):
        """Extract text content from a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            contents = [text.strip() for text in f.readlines() if text.strip()]
        return contents

    @staticmethod
    def extract_text_from_docx(file_path: str):
        """Extract text content from a DOCX file."""
        import docx
        document = docx.Document(file_path)
        contents = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return contents

    @staticmethod
    def extract_text_from_markdown(file_path: str):
        """Extract text content from a Markdown file."""
        import markdown
        from bs4 import BeautifulSoup
        with open(file_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        html = markdown.markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        contents = [text.strip() for text in soup.get_text().splitlines() if text.strip()]
        return contents